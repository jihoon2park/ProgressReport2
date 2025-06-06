#!/usr/bin/env python3
"""
내부 서버 배포 가이드 문서 변환 스크립트
마크다운 파일들을 워드(.docx)와 PDF 파일로 변환합니다.
"""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown
import weasyprint

def read_markdown_file(filepath):
    """마크다운 파일을 읽습니다."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        print(f"❌ 파일을 찾을 수 없습니다: {filepath}")
        return None
    except Exception as e:
        print(f"❌ 파일 읽기 오류: {filepath} - {e}")
        return None

def markdown_to_docx(markdown_content, title, output_file):
    """마크다운 내용을 워드 문서로 변환합니다."""
    doc = Document()
    
    # 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)
    
    # 제목
    title_paragraph = doc.add_heading(title, 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 생성 날짜
    date_paragraph = doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # 마크다운 내용 파싱
    lines = markdown_content.split('\n')
    in_code_block = False
    code_lang = ''
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 코드 블록 처리
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line[3:].strip()
                i += 1
                continue
            else:
                in_code_block = False
                code_lang = ''
                i += 1
                continue
        
        if in_code_block:
            # 코드 블록 내용
            code_para = doc.add_paragraph(line)
            code_para.style = 'Intense Quote'
            run = code_para.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        else:
            # 일반 텍스트 처리
            if line.startswith('# '):
                # H1 헤딩
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # H2 헤딩
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # H3 헤딩
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                # H4 헤딩
                doc.add_heading(line[5:], level=4)
            elif line.startswith('- '):
                # 불릿 리스트
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.', line):
                # 번호 리스트
                doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
            elif line.startswith('> '):
                # 인용문
                para = doc.add_paragraph(line[2:])
                para.style = 'Quote'
            elif line.strip() == '':
                # 빈 줄
                doc.add_paragraph('')
            else:
                # 일반 텍스트
                if line.strip():
                    # **굵게**, *기울임*, `코드` 처리
                    para = doc.add_paragraph()
                    process_inline_formatting(para, line)
        
        i += 1
    
    # 문서 저장
    doc.save(output_file)
    print(f"✅ 워드 문서 생성 완료: {output_file}")

def process_inline_formatting(paragraph, text):
    """인라인 포맷팅을 처리합니다."""
    # 단순화된 버전 - 복잡한 마크다운 구문은 생략
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 굵게 제거 (워드에서는 따로 처리)
    text = re.sub(r'`([^`]*)`', r'\1', text)      # 코드 마크 제거
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # 링크에서 텍스트만 추출
    
    paragraph.add_run(text)

def markdown_to_pdf(markdown_content, title, output_file):
    """마크다운 내용을 PDF로 변환합니다."""
    # HTML로 변환
    html_content = markdown.markdown(
        markdown_content,
        extensions=['codehilite', 'fenced_code', 'tables', 'toc']
    )
    
    # CSS 스타일
    css_style = """
    <style>
        body {
            font-family: 'Malgun Gothic', Arial, sans-serif;
            line-height: 1.6;
            margin: 2cm;
            font-size: 12pt;
        }
        h1, h2, h3, h4, h5, h6 {
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }
        h1 {
            text-align: center;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }
        h2 {
            border-bottom: 2px solid #e74c3c;
            padding-bottom: 5px;
        }
        h3 {
            color: #e67e22;
        }
        code {
            background-color: #f8f9fa;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: 'Consolas', monospace;
            font-size: 10pt;
        }
        pre {
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            border-left: 4px solid #3498db;
            overflow-x: auto;
            font-family: 'Consolas', monospace;
            font-size: 10pt;
        }
        blockquote {
            border-left: 4px solid #bdc3c7;
            margin: 0;
            padding-left: 15px;
            font-style: italic;
            color: #7f8c8d;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }
        table, th, td {
            border: 1px solid #ddd;
        }
        th, td {
            padding: 8px;
            text-align: left;
        }
        th {
            background-color: #f2f2f2;
        }
        ul, ol {
            margin: 0.5em 0;
            padding-left: 2em;
        }
        .header-info {
            text-align: right;
            color: #7f8c8d;
            font-size: 10pt;
            margin-bottom: 2em;
        }
    </style>
    """
    
    # 완전한 HTML 문서 생성
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{title}</title>
        {css_style}
    </head>
    <body>
        <h1>{title}</h1>
        <div class="header-info">생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}</div>
        {html_content}
    </body>
    </html>
    """
    
    # PDF 생성
    try:
        weasyprint.HTML(string=full_html).write_pdf(output_file)
        print(f"✅ PDF 문서 생성 완료: {output_file}")
    except Exception as e:
        print(f"❌ PDF 생성 오류: {e}")

def main():
    """메인 함수"""
    print("📄 내부 서버 배포 가이드 문서 변환 시작...")
    
    # 변환할 문서 리스트
    documents = [
        {
            'file': 'INTERNAL_DEPLOYMENT.md',
            'title': '🏢 회사 내부 서버 배포 가이드',
            'output_prefix': '내부서버_배포가이드'
        },
        {
            'file': 'WINDOWS_TO_INTERNAL.md', 
            'title': '💻 윈도우에서 내부 서버로 배포하기',
            'output_prefix': '윈도우_배포방법'
        }
    ]
    
    # 출력 디렉토리 생성
    output_dir = 'documents'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"📁 출력 디렉토리 생성: {output_dir}")
    
    # 각 문서 변환
    for doc_info in documents:
        markdown_content = read_markdown_file(doc_info['file'])
        if markdown_content:
            # 워드 문서 생성
            docx_file = os.path.join(output_dir, f"{doc_info['output_prefix']}.docx")
            markdown_to_docx(markdown_content, doc_info['title'], docx_file)
            
            # PDF 문서 생성
            pdf_file = os.path.join(output_dir, f"{doc_info['output_prefix']}.pdf")
            markdown_to_pdf(markdown_content, doc_info['title'], pdf_file)
        else:
            print(f"⚠️ 파일을 건너뜀: {doc_info['file']}")
    
    # 통합 문서 생성
    print("\n📚 통합 문서 생성 중...")
    create_combined_document()
    
    print("\n✅ 모든 문서 변환 완료!")
    print(f"📁 출력 위치: {os.path.abspath(output_dir)}")
    print("\n생성된 파일들:")
    for file in os.listdir(output_dir):
        print(f"  📄 {file}")

def create_combined_document():
    """모든 가이드를 하나로 합친 통합 문서를 생성합니다."""
    # 모든 마크다운 내용 합치기
    combined_content = "# 🏢 회사 내부 서버 배포 통합 가이드\n\n"
    combined_content += f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
    combined_content += "---\n\n"
    
    # 목차 추가
    combined_content += "## 📋 목차\n\n"
    combined_content += "1. [내부 서버 배포 가이드](#내부-서버-배포-가이드)\n"
    combined_content += "2. [윈도우에서 배포하기](#윈도우에서-배포하기)\n"
    combined_content += "3. [설정 파일들](#설정-파일들)\n\n"
    combined_content += "---\n\n"
    
    # 내부 서버 배포 가이드
    internal_content = read_markdown_file('INTERNAL_DEPLOYMENT.md')
    if internal_content:
        combined_content += "## 내부 서버 배포 가이드\n\n"
        combined_content += internal_content + "\n\n---\n\n"
    
    # 윈도우 배포 방법
    windows_content = read_markdown_file('WINDOWS_TO_INTERNAL.md')
    if windows_content:
        combined_content += "## 윈도우에서 배포하기\n\n"
        combined_content += windows_content + "\n\n---\n\n"
    
    # 설정 파일 내용 추가
    combined_content += "## 설정 파일들\n\n"
    
    # Nginx 설정
    nginx_content = read_markdown_file('nginx_internal.conf')
    if nginx_content:
        combined_content += "### nginx_internal.conf\n\n"
        combined_content += "```nginx\n"
        combined_content += nginx_content
        combined_content += "\n```\n\n"
    
    # 배포 스크립트
    deploy_content = read_markdown_file('deploy_internal.sh')
    if deploy_content:
        combined_content += "### deploy_internal.sh\n\n"
        combined_content += "```bash\n"
        combined_content += deploy_content
        combined_content += "\n```\n\n"
    
    # 통합 문서 생성
    output_dir = 'documents'
    
    # 워드 문서
    docx_file = os.path.join(output_dir, "내부서버_배포_통합가이드.docx")
    markdown_to_docx(combined_content, "🏢 회사 내부 서버 배포 통합 가이드", docx_file)
    
    # PDF 문서
    pdf_file = os.path.join(output_dir, "내부서버_배포_통합가이드.pdf")
    markdown_to_pdf(combined_content, "🏢 회사 내부 서버 배포 통합 가이드", pdf_file)

if __name__ == "__main__":
    main() 