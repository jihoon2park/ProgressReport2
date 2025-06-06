#!/usr/bin/env python3
"""
내부 서버 배포 가이드 문서 변환 스크립트 (윈도우 최적화 버전)
마크다운 파일들을 워드(.docx)와 PDF 파일로 변환합니다.
"""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import markdown2

def read_file(filepath):
    """파일을 읽습니다."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        print(f"❌ 파일을 찾을 수 없습니다: {filepath}")
        return None
    except Exception as e:
        print(f"❌ 파일 읽기 오류: {filepath} - {e}")
        return None

def markdown_to_docx(markdown_content, title, output_file):
    """마크다운 내용을 워드 문서로 변환합니다."""
    doc = Document()
    
    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)
    
    # 제목 추가
    title_paragraph = doc.add_heading(title, 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 생성 날짜
    date_paragraph = doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 페이지 구분
    doc.add_page_break()
    
    # 마크다운 내용 파싱
    lines = markdown_content.split('\n')
    in_code_block = False
    code_lines = []
    
    for line in lines:
        line = line.rstrip()
        
        # 코드 블록 처리
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                continue
            else:
                # 코드 블록 종료 - 코드 내용 추가
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_para.style = 'Intense Quote'
                    run = code_para.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                in_code_block = False
                code_lines = []
                continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # 헤딩 처리
        if line.startswith('# ') and not line.startswith('## '):
            heading_text = line[2:].strip()
            # 이모지 제거
            heading_text = re.sub(r'[^\w\s가-힣]', '', heading_text).strip()
            doc.add_heading(heading_text, level=1)
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            heading_text = re.sub(r'[^\w\s가-힣]', '', heading_text).strip()
            doc.add_heading(heading_text, level=2)
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            heading_text = re.sub(r'[^\w\s가-힣]', '', heading_text).strip()
            doc.add_heading(heading_text, level=3)
        elif line.startswith('#### '):
            heading_text = line[5:].strip()
            heading_text = re.sub(r'[^\w\s가-힣]', '', heading_text).strip()
            doc.add_heading(heading_text, level=4)
        # 리스트 처리
        elif line.startswith('- '):
            list_text = line[2:].strip()
            list_text = clean_markdown_text(list_text)
            doc.add_paragraph(list_text, style='List Bullet')
        elif re.match(r'^\d+\.', line):
            list_text = re.sub(r'^\d+\.\s*', '', line).strip()
            list_text = clean_markdown_text(list_text)
            doc.add_paragraph(list_text, style='List Number')
        # 인용문 처리
        elif line.startswith('> '):
            quote_text = line[2:].strip()
            quote_text = clean_markdown_text(quote_text)
            para = doc.add_paragraph(quote_text)
            para.style = 'Quote'
        # 빈 줄
        elif line.strip() == '':
            doc.add_paragraph('')
        # 구분선
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
        # 일반 텍스트
        else:
            if line.strip():
                text = clean_markdown_text(line)
                if text:
                    doc.add_paragraph(text)
    
    # 문서 저장
    try:
        doc.save(output_file)
        print(f"✅ 워드 문서 생성 완료: {output_file}")
        return True
    except Exception as e:
        print(f"❌ 워드 문서 생성 오류: {e}")
        return False

def clean_markdown_text(text):
    """마크다운 서식을 제거하고 텍스트만 추출합니다."""
    # 굵게, 기울임 제거
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # 코드 마크 제거
    text = re.sub(r'`([^`]*)`', r'\1', text)
    # 링크에서 텍스트만 추출
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    # HTML 태그 제거
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()

def markdown_to_pdf(markdown_content, title, output_file):
    """마크다운 내용을 PDF로 변환합니다."""
    try:
        # PDF 문서 생성
        doc = SimpleDocTemplate(
            output_file,
            pagesize=A4,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        
        # 스타일 정의
        styles = getSampleStyleSheet()
        
        # 제목 스타일
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor='#2c3e50'
        )
        
        # 헤딩 스타일들
        h1_style = ParagraphStyle(
            'CustomH1',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=20,
            textColor='#2c3e50'
        )
        
        h2_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=16,
            textColor='#34495e'
        )
        
        h3_style = ParagraphStyle(
            'CustomH3',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=8,
            spaceBefore=12,
            textColor='#7f8c8d'
        )
        
        # 일반 텍스트 스타일
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            leading=14
        )
        
        # 코드 스타일
        code_style = ParagraphStyle(
            'CustomCode',
            parent=styles['Code'],
            fontSize=9,
            leftIndent=20,
            backgroundColor='#f8f9fa',
            borderColor='#e9ecef',
            borderWidth=1,
            borderPadding=5
        )
        
        # 컨텐츠 리스트
        story = []
        
        # 제목과 날짜
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 
                             ParagraphStyle('Date', parent=styles['Normal'], alignment=TA_RIGHT, fontSize=9)))
        story.append(Spacer(1, 20))
        
        # 마크다운 내용 파싱
        lines = markdown_content.split('\n')
        in_code_block = False
        code_lines = []
        
        for line in lines:
            line = line.rstrip()
            
            # 코드 블록 처리
            if line.startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_lines = []
                    continue
                else:
                    # 코드 블록 종료
                    if code_lines:
                        code_text = '\n'.join(code_lines)
                        story.append(Paragraph(code_text.replace('<', '&lt;').replace('>', '&gt;'), code_style))
                        story.append(Spacer(1, 6))
                    in_code_block = False
                    code_lines = []
                    continue
            
            if in_code_block:
                code_lines.append(line)
                continue
            
            # 헤딩 처리
            if line.startswith('# ') and not line.startswith('## '):
                heading_text = clean_markdown_text(line[2:].strip())
                story.append(Paragraph(heading_text, h1_style))
            elif line.startswith('## '):
                heading_text = clean_markdown_text(line[3:].strip())
                story.append(Paragraph(heading_text, h2_style))
            elif line.startswith('### '):
                heading_text = clean_markdown_text(line[4:].strip())
                story.append(Paragraph(heading_text, h3_style))
            # 구분선
            elif line.strip() == '---':
                story.append(Spacer(1, 10))
                story.append(Paragraph('─' * 80, normal_style))
                story.append(Spacer(1, 10))
            # 빈 줄
            elif line.strip() == '':
                story.append(Spacer(1, 6))
            # 일반 텍스트
            else:
                if line.strip():
                    text = clean_markdown_text(line)
                    if text:
                        story.append(Paragraph(text, normal_style))
        
        # PDF 생성
        doc.build(story)
        print(f"✅ PDF 문서 생성 완료: {output_file}")
        return True
        
    except Exception as e:
        print(f"❌ PDF 생성 오류: {e}")
        return False

def main():
    """메인 함수"""
    print("📄 내부 서버 배포 가이드 문서 변환 시작...")
    
    # 출력 디렉토리 생성
    output_dir = 'documents'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"📁 출력 디렉토리 생성: {output_dir}")
    
    # 변환할 문서 리스트
    documents = [
        {
            'file': 'INTERNAL_DEPLOYMENT.md',
            'title': '회사 내부 서버 배포 가이드',
            'output_prefix': '내부서버_배포가이드'
        },
        {
            'file': 'WINDOWS_TO_INTERNAL.md', 
            'title': '윈도우에서 내부 서버로 배포하기',
            'output_prefix': '윈도우_배포방법'
        }
    ]
    
    success_count = 0
    
    # 각 문서 변환
    for doc_info in documents:
        print(f"\n📄 처리 중: {doc_info['file']}")
        
        markdown_content = read_file(doc_info['file'])
        if markdown_content:
            # 워드 문서 생성
            docx_file = os.path.join(output_dir, f"{doc_info['output_prefix']}.docx")
            if markdown_to_docx(markdown_content, doc_info['title'], docx_file):
                success_count += 1
            
            # PDF 문서 생성
            pdf_file = os.path.join(output_dir, f"{doc_info['output_prefix']}.pdf")
            if markdown_to_pdf(markdown_content, doc_info['title'], pdf_file):
                success_count += 1
        else:
            print(f"⚠️ 파일을 건너뜀: {doc_info['file']}")
    
    # 통합 문서 생성
    print("\n📚 통합 문서 생성 중...")
    if create_combined_document(output_dir):
        success_count += 2
    
    print(f"\n✅ 문서 변환 완료! ({success_count}개 파일 생성)")
    print(f"📁 출력 위치: {os.path.abspath(output_dir)}")
    
    # 생성된 파일 목록 출력
    if os.path.exists(output_dir):
        files = [f for f in os.listdir(output_dir) if f.endswith(('.docx', '.pdf'))]
        if files:
            print("\n📄 생성된 파일들:")
            for file in sorted(files):
                file_size = os.path.getsize(os.path.join(output_dir, file))
                size_kb = file_size / 1024
                print(f"  📄 {file} ({size_kb:.1f}KB)")

def create_combined_document(output_dir):
    """모든 가이드를 하나로 합친 통합 문서를 생성합니다."""
    try:
        # 모든 마크다운 내용 합치기
        combined_content = "# 회사 내부 서버 배포 통합 가이드\n\n"
        combined_content += f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
        combined_content += "---\n\n"
        
        # 목차 추가
        combined_content += "## 목차\n\n"
        combined_content += "1. 내부 서버 배포 가이드\n"
        combined_content += "2. 윈도우에서 배포하기\n"
        combined_content += "3. 설정 파일들\n\n"
        combined_content += "---\n\n"
        
        # 내부 서버 배포 가이드
        internal_content = read_file('INTERNAL_DEPLOYMENT.md')
        if internal_content:
            combined_content += "## 내부 서버 배포 가이드\n\n"
            combined_content += internal_content + "\n\n---\n\n"
        
        # 윈도우 배포 방법
        windows_content = read_file('WINDOWS_TO_INTERNAL.md')
        if windows_content:
            combined_content += "## 윈도우에서 배포하기\n\n"
            combined_content += windows_content + "\n\n---\n\n"
        
        # 설정 파일 내용 추가
        combined_content += "## 주요 설정 파일들\n\n"
        
        # Nginx 설정
        nginx_content = read_file('nginx_internal.conf')
        if nginx_content:
            combined_content += "### nginx_internal.conf\n\n"
            combined_content += "```\n"
            combined_content += nginx_content[:2000] + "...\n```\n\n"  # 너무 길면 축약
        
        # 통합 문서 생성
        # 워드 문서
        docx_file = os.path.join(output_dir, "내부서버_배포_통합가이드.docx")
        docx_success = markdown_to_docx(combined_content, "회사 내부 서버 배포 통합 가이드", docx_file)
        
        # PDF 문서
        pdf_file = os.path.join(output_dir, "내부서버_배포_통합가이드.pdf")
        pdf_success = markdown_to_pdf(combined_content, "회사 내부 서버 배포 통합 가이드", pdf_file)
        
        return docx_success and pdf_success
        
    except Exception as e:
        print(f"❌ 통합 문서 생성 오류: {e}")
        return False

if __name__ == "__main__":
    main() 